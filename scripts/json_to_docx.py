#!/usr/bin/env python3
"""
JSON to DOCX Converter for Bootcamp Workshop
==============================================
Convert bootcamp JSON data to professional DOCX document.
"""

import json
import sys
from typing import Dict, Any, List
# Fix encoding for Windows
if sys.stdout:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except:
        pass

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class BootcampToDocx:
    """Convert Bootcamp JSON to DOCX document."""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup document styles."""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
    
    def _add_heading(self, text: str, level: int = 1):
        """Add formatted heading."""
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return heading
    
    def _add_paragraph(self, text: str, bold: bool = False, italic: bool = False):
        """Add formatted paragraph."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return p
    
    def _add_table_border(self, table):
        """Add borders to table."""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    def _add_cover_page(self, data: Dict[str, Any]):
        """Add cover page."""
        # Title
        title = self.doc.add_heading(data['identitas']['nama'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle_text = data.get('deskripsiSingkat', '')
        subtitle = self.doc.add_paragraph(subtitle_text)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if subtitle.runs:
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.size = Pt(14)
            subtitle_run.italic = True
        
        self.doc.add_paragraph()  # Space
        
        # Bootcamp info table
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Kode', data['identitas']['kode']),
            ('Durasi', f"{data['identitas']['durasi']} Minggu"),
            ('Level', data['identitas']['level']),
            ('Tipe', data['identitas']['tipe']),
            ('Kapasitas', f"{data['identitas']['kapasitas']} Peserta")
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        self.doc.add_page_break()
    
    def _add_description(self, data: Dict[str, Any]):
        """Add description section."""
        self._add_heading('Deskripsi Bootcamp', level=1)
        self._add_paragraph(data.get('deskripsi', ''))
        self.doc.add_paragraph()
    
    def _add_target_peserta(self, data: Dict[str, Any]):
        """Add target participants section."""
        self._add_heading('Target Peserta', level=1)
        
        target = data.get('targetPeserta', {})
        
        # Deskripsi umum
        self._add_paragraph(target.get('deskripsi', ''))
        self.doc.add_paragraph()
        
        # Latar Belakang
        self._add_heading('Latar Belakang Yang Cocok:', level=2)
        for item in target.get('latar_belakang', []):
            self.doc.add_paragraph(item, style='List Bullet')
        
        # Prasyarat Teknis
        self._add_heading('Prasyarat Teknis:', level=2)
        for item in target.get('prasyarat_teknis', []):
            self.doc.add_paragraph(item, style='List Bullet')
        
        # Prasyarat Soft Skills
        self._add_heading('Soft Skills Yang Diharapkan:', level=2)
        for item in target.get('prasyarat_soft_skill', []):
            self.doc.add_paragraph(item, style='List Bullet')
        
        self.doc.add_paragraph()
    
    def _add_learning_outcomes(self, data: Dict[str, Any]):
        """Add learning outcomes section."""
        self._add_heading('Learning Outcomes', level=1)
        
        los = data.get('learningOutcomes', [])
        if not los:
            return
        
        # Group by category
        categories = {}
        for lo in los:
            cat = lo.get('kategori', 'Other')
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(lo)
        
        for category, items in categories.items():
            self._add_heading(f'{category} Skills:', level=2)
            for lo in items:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(f"{lo['kode']}: ").bold = True
                p.add_run(lo['pernyataan'])
        
        self.doc.add_paragraph()
    
    def _add_weekly_schedule(self, data: Dict[str, Any]):
        """Add weekly schedule section."""
        self._add_heading('Jadwal Pembelajaran Mingguan', level=1)
        
        minggu_list = data.get('minggu', [])
        
        for minggu in minggu_list:
            week_num = minggu['mingguKe']
            
            # Week header
            self._add_heading(f"Minggu {week_num}: {minggu['tema']}", level=2)
            
            # Learning outcomes for this week
            p = self.doc.add_paragraph()
            p.add_run('Learning Outcomes: ').bold = True
            p.add_run(', '.join(minggu.get('learningOutcomes', [])))
            
            # Materi Pokok
            if 'materiPokok' in minggu:
                self.doc.add_paragraph()
                self._add_paragraph('Materi Pokok:', bold=True)
                for materi in minggu['materiPokok']:
                    self.doc.add_paragraph(materi, style='List Bullet')
            
            # Metode Pembelajaran
            if 'metodePembelajaran' in minggu:
                metode = minggu['metodePembelajaran']
                self.doc.add_paragraph()
                p = self.doc.add_paragraph()
                p.add_run('Metode: ').bold = True
                p.add_run(metode['metode'])
                
                self._add_paragraph(f"Deskripsi: {metode['deskripsi']}")
                self._add_paragraph(f"Aktivitas: {metode['aktivitas']}")
            
            # Waktu
            if 'waktu' in minggu:
                p = self.doc.add_paragraph()
                p.add_run('Alokasi Waktu: ').bold = True
                p.add_run(minggu['waktu'])
            
            # Project (if any)
            if 'project' in minggu:
                project = minggu['project']
                self.doc.add_paragraph()
                self._add_heading(f"📦 Project: {project['nama']}", level=3)
                self._add_paragraph(project['deskripsi'])
                
                self._add_paragraph('Deliverables:', bold=True)
                for deliverable in project.get('deliverables', []):
                    self.doc.add_paragraph(deliverable, style='List Bullet')
                
                p = self.doc.add_paragraph()
                p.add_run('Teknologi: ').bold = True
                p.add_run(', '.join(project.get('teknologi', [])))
            
            # Pengalaman Belajar
            if 'pengalamanBelajar' in minggu:
                self.doc.add_paragraph()
                p = self.doc.add_paragraph()
                p.add_run('Pengalaman Belajar: ').bold = True
                p.add_run(minggu['pengalamanBelajar'])
            
            # Penilaian
            if 'penilaian' in minggu:
                penilaian = minggu['penilaian']
                self.doc.add_paragraph()
                p = self.doc.add_paragraph()
                p.add_run('Kriteria Penilaian: ').bold = True
                p.add_run(penilaian['kriteria'])
                
                p = self.doc.add_paragraph()
                p.add_run('Bobot: ').bold = True
                p.add_run(f"{penilaian['bobot']}%")
            
            self.doc.add_paragraph()  # Space between weeks
    
    def _add_assessment(self, data: Dict[str, Any]):
        """Add assessment section."""
        self._add_heading('Komponen Penilaian', level=1)
        
        assessments = data.get('assessment', [])
        
        # Create table
        table = self.doc.add_table(rows=len(assessments) + 1, cols=3)
        self._add_table_border(table)
        
        # Header
        header_cells = table.rows[0].cells
        headers = ['Komponen', 'Bobot', 'Metode & Kriteria']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Data
        for i, assessment in enumerate(assessments, 1):
            row = table.rows[i]
            row.cells[0].text = assessment['nama']
            row.cells[1].text = f"{assessment['bobot']}%"
            
            # Metode & Kriteria
            criteria_text = f"Metode: {assessment['metode']}\n\n"
            criteria_text += f"Deskripsi: {assessment['deskripsi']}\n\n"
            criteria_text += "Kriteria:\n"
            for criterion in assessment.get('kriteria', []):
                criteria_text += f"• {criterion}\n"
            row.cells[2].text = criteria_text
        
        self.doc.add_paragraph()
    
    def _add_instructors(self, data: Dict[str, Any]):
        """Add instructors section."""
        self._add_heading('Tim Instruktur', level=1)
        
        instructors = data.get('instruktur', [])
        
        for instructor in instructors:
            p = self.doc.add_paragraph()
            p.add_run(f"{instructor['nama']} - {instructor['peran']}").bold = True
            
            self._add_paragraph(f"Keahlian: {', '.join(instructor['expertise'])}")
            
            if 'kontak' in instructor:
                self._add_paragraph(f"Kontak: {instructor['kontak']}")
            
            self.doc.add_paragraph()
    
    def _add_tools_resources(self, data: Dict[str, Any]):
        """Add tools and resources section."""
        self._add_heading('Tools & Resources', level=1)
        
        tools = data.get('toolsResources', {})
        
        # Software
        if 'software' in tools:
            self._add_heading('Software:', level=2)
            for item in tools['software']:
                self.doc.add_paragraph(item, style='List Bullet')
        
        # Platform
        if 'platform' in tools:
            self._add_heading('Platform:', level=2)
            for item in tools['platform']:
                self.doc.add_paragraph(item, style='List Bullet')
        
        # Hardware
        if 'hardware' in tools:
            self._add_heading('Hardware Requirements:', level=2)
            for item in tools['hardware']:
                self.doc.add_paragraph(item, style='List Bullet')
        
        # Akun
        if 'akun_diperlukan' in tools:
            self._add_heading('Akun Yang Perlu Dibuat:', level=2)
            for item in tools['akun_diperlukan']:
                self.doc.add_paragraph(item, style='List Bullet')
        
        self.doc.add_paragraph()
    
    def _add_certification(self, data: Dict[str, Any]):
        """Add certification section."""
        self._add_heading('Sertifikasi', level=1)
        
        cert = data.get('sertifikasi', {})
        
        p = self.doc.add_paragraph()
        p.add_run('Nama Sertifikat: ').bold = True
        p.add_run(cert.get('nama', ''))
        
        p = self.doc.add_paragraph()
        p.add_run('Nilai Minimal Kelulusan: ').bold = True
        p.add_run(f"{cert.get('nilai_minimal', 0)}%")
        
        self._add_heading('Syarat Kelulusan:', level=2)
        for syarat in cert.get('syarat_kelulusan', []):
            self.doc.add_paragraph(syarat, style='List Bullet')
        
        self._add_heading('Benefit:', level=2)
        for benefit in cert.get('benefit', []):
            self.doc.add_paragraph(benefit, style='List Bullet')
        
        self.doc.add_paragraph()
    
    def _add_references(self, data: Dict[str, Any]):
        """Add references section."""
        self._add_heading('Referensi', level=1)
        
        references = data.get('referensi', [])
        for i, ref in enumerate(references, 1):
            self.doc.add_paragraph(f"{i}. {ref}")
        
        self.doc.add_paragraph()
    
    def _add_facilities(self, data: Dict[str, Any]):
        """Add facilities section if exists."""
        if 'fasilitas' not in data:
            return
        
        self._add_heading('Fasilitas', level=1)
        for fasilitas in data['fasilitas']:
            self.doc.add_paragraph(fasilitas, style='List Bullet')
        
        self.doc.add_paragraph()
    
    def _add_investment(self, data: Dict[str, Any]):
        """Add investment/pricing section if exists."""
        if 'investasi' not in data:
            return
        
        self._add_heading('Investasi', level=1)
        
        inv = data['investasi']
        
        p = self.doc.add_paragraph()
        p.add_run('Biaya Normal: ').bold = True
        p.add_run(f"Rp {inv.get('biaya', 0):,}")
        
        if 'early_bird' in inv:
            p = self.doc.add_paragraph()
            p.add_run('Harga Early Bird: ').bold = True
            p.add_run(f"Rp {inv['early_bird']:,}")
        
        if inv.get('cicilan'):
            self._add_paragraph('✅ Tersedia opsi cicilan')
        
        if inv.get('beasiswa'):
            self._add_paragraph('✅ Tersedia program beasiswa')
        
        self.doc.add_paragraph()
    
    def convert(self, json_file: str, output_file: str):
        """
        Convert JSON file to DOCX.
        
        Args:
            json_file: Path to input JSON file
            output_file: Path to output DOCX file
        """
        print(f"📄 Converting {json_file} to {output_file}")
        
        # Load JSON data
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Build document
        self._add_cover_page(data)
        self._add_description(data)
        self._add_target_peserta(data)
        self._add_learning_outcomes(data)
        self._add_weekly_schedule(data)
        self._add_assessment(data)
        self._add_instructors(data)
        self._add_tools_resources(data)
        self._add_certification(data)
        self._add_facilities(data)
        self._add_investment(data)
        self._add_references(data)
        
        # Save document
        self.doc.save(output_file)
        print(f"✅ DOCX saved to: {output_file}")


# Standalone usage
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Convert Bootcamp JSON to DOCX")
    parser.add_argument("--input", "-i", default="bootcamp_generated.json", help="Input JSON file")
    parser.add_argument("--output", "-o", default="bootcamp_curriculum.docx", help="Output DOCX file")
    
    args = parser.parse_args()
    
    print("DOCX Converter Started...")
    print("=" * 60)
    
    converter = BootcampToDocx()
    converter.convert(args.input, args.output)
